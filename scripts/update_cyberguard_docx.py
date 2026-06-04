#!/usr/bin/env python3
"""
Transform Project 404 documentation template into CyberGuard AI documentation.

Preserves heading hierarchy, numbering style, tables, and captions where possible.
Rewrites body content from scripts/cyberguard_docx_content.py based on the actual codebase.

Usage (from repository root):
    python scripts/update_cyberguard_docx.py
    python scripts/update_cyberguard_docx.py --input "Project 404 Final Version 1.docx" --output "docs/CyberGuard AI - Technical Documentation.docx"
"""

from __future__ import annotations

import argparse
import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent))

from cyberguard_docx_content import (
    API_ENDPOINTS_TABLE,
    EVENT_MODEL_TABLE,
    GLOBAL_REPLACEMENTS,
    HEADING_RENAMES,
    PROJECT_NAME,
    SECTION_BODIES,
    TECH_STACK_TABLE,
)

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = REPO_ROOT / "Project 404 Final Version 1.docx"
DEFAULT_OUTPUT = REPO_ROOT / "docs" / "CyberGuard AI - Technical Documentation.docx"


def _is_heading(paragraph: Paragraph) -> bool:
    name = paragraph.style.name if paragraph.style else ""
    return bool(name and name.startswith("Heading"))


def _heading_level(paragraph: Paragraph) -> int:
    name = paragraph.style.name if paragraph.style else ""
    m = re.search(r"Heading (\d+)", name)
    return int(m.group(1)) if m else 0


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _clear_runs_keep_style(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert_paragraph_after(
    reference: Paragraph, text: str, style_name: str | None = None
) -> Paragraph:
    new_p = deepcopy(reference._element)
    reference._element.addnext(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    _set_paragraph_text(paragraph, text)
    if style_name:
        try:
            paragraph.style = style_name
        except KeyError:
            paragraph.style = reference.style
    return paragraph


def _replace_section_body(
    doc: Document,
    start_idx: int,
    end_idx: int,
    new_paragraphs: list[str],
    body_style: str = "Normal",
) -> None:
    """Replace non-heading paragraphs in [start_idx, end_idx) with new content."""
    if start_idx >= end_idx:
        return
    ref = doc.paragraphs[start_idx]
    for idx in range(end_idx - 1, start_idx, -1):
        if not _is_heading(doc.paragraphs[idx]):
            _delete_paragraph(doc.paragraphs[idx])
    anchor = doc.paragraphs[start_idx]
    prev = anchor
    for text in new_paragraphs:
        prev = _insert_paragraph_after(prev, text, body_style)


def _build_heading_index(doc: Document) -> list[tuple[int, str, int]]:
    index: list[tuple[int, str, int]] = []
    for i, p in enumerate(doc.paragraphs):
        if _is_heading(p):
            index.append((i, p.text.strip(), _heading_level(p)))
    return index


def _find_section_range(
    heading_index: list[tuple[int, str, int]], key: str, level: int = 2
) -> tuple[int, int] | None:
    """Return (start_para_idx, end_para_idx) for body under a heading key."""
    start = None
    start_level = level
    for i, (idx, text, lvl) in enumerate(heading_index):
        if key in text and lvl == level and start is None:
            start = idx
            start_level = lvl
            continue
        if start is not None and lvl <= start_level and idx > start:
            return start, idx
    if start is not None:
        return start, 10_000
    return None


def _apply_heading_renames(doc: Document) -> int:
    count = 0
    for p in doc.paragraphs:
        if not _is_heading(p):
            continue
        text = p.text.strip()
        for old, new in HEADING_RENAMES.items():
            if text == old or old in text:
                _set_paragraph_text(p, new if text == old else text.replace(old, new))
                count += 1
                break
    return count


def _apply_section_bodies(doc: Document) -> int:
    heading_index = _build_heading_index(doc)
    applied = 0
    for key, body in SECTION_BODIES.items():
        rng = _find_section_range(heading_index, key, level=2)
        if not rng:
            continue
        start, end = rng
        end = min(end, len(doc.paragraphs))
        _replace_section_body(doc, start, end, body)
        applied += 1
        heading_index = _build_heading_index(doc)
    return applied


def _apply_global_replacements(doc: Document) -> int:
    count = 0
    for p in doc.paragraphs:
        original = p.text
        updated = original
        for old, new in GLOBAL_REPLACEMENTS:
            if old in updated:
                updated = updated.replace(old, new)
        if updated != original:
            _set_paragraph_text(p, updated)
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original = p.text
                    updated = original
                    for old, new in GLOBAL_REPLACEMENTS:
                        if old in updated:
                            updated = updated.replace(old, new)
                    if updated != original:
                        _set_paragraph_text(p, updated)
                        count += 1
    return count


def _fill_table(table: Table, rows: list[list[str]], header: bool = True) -> None:
    needed = len(rows)
    while len(table.rows) < needed:
        table.add_row()
    while len(table.rows) > needed:
        table._tbl.remove(table.rows[-1]._tr)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        while len(row.cells) < len(row_data):
            pass
        for ci, val in enumerate(row_data):
            if ci < len(row.cells):
                _set_paragraph_text(row.cells[ci].paragraphs[0], val)


def _update_stakeholder_table(doc: Document) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if "Stakeholder" in " ".join(header) and "Influence" in " ".join(header):
            rows = [
                header,
                [
                    "Student / Developer Team",
                    "Design, implement, and test CyberGuard AI",
                    "High",
                    "High",
                    "Owns codebase and documentation.",
                ],
                [
                    "Academic Supervisor / Client",
                    "Review outcomes and documentation",
                    "Medium",
                    "Medium",
                    "Approves scope and deliverables.",
                ],
                [
                    "Security Analyst (end user)",
                    "Triage alerts and incidents",
                    "High",
                    "High",
                    "Primary operational user.",
                ],
                [
                    "Platform Administrator",
                    "Deploy Docker Compose stack",
                    "Medium",
                    "Medium",
                    "Manages Postgres, Redis, secrets.",
                ],
            ]
            _fill_table(table, rows, header=True)


def _update_use_case_tables(doc: Document) -> None:
    use_cases = [
        ("UC-CG-001", "Ingest Security Events", "Analyst submits events via API or file upload."),
        ("UC-CG-002", "View Dashboard", "Analyst reviews KPIs and charts for selected window."),
        ("UC-CG-003", "Triage Alert", "Analyst acknowledges or resolves an alert."),
        ("UC-CG-004", "Investigate Incident", "Analyst opens incident, links events, runs analysis."),
        ("UC-CG-005", "Generate PDF Report", "Analyst queues and downloads incident PDF."),
        ("UC-CG-006", "Run Cyber Range Scenario", "Analyst runs approved local attack scenario."),
        ("UC-CG-007", "Configure Digest Schedule", "User sets scheduled PDF digest and email."),
    ]
    uc_idx = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        first = table.rows[0].cells[0].text.strip()
        if first == "ID" and "UC-" in table.rows[0].cells[1].text:
            if uc_idx < len(use_cases):
                uid, title, desc = use_cases[uc_idx]
                data = [
                    ["ID", uid],
                    ["Title", title],
                    ["Description", desc],
                    ["Actor", "Authenticated analyst"],
                    ["Preconditions", "User logged in; API available"],
                    ["Postconditions", "Events/alerts/incidents updated per workflow"],
                    ["Priority", "High"],
                ]
                for ri in range(min(len(data), len(table.rows))):
                    _set_paragraph_text(table.rows[ri].cells[0].paragraphs[0], data[ri][0])
                    if len(table.rows[ri].cells) > 1:
                        _set_paragraph_text(
                            table.rows[ri].cells[1].paragraphs[0], data[ri][1]
                        )
                uc_idx += 1


def _insert_api_section_after_tools(doc: Document) -> None:
    """Add API documentation subsection after 5.1 if not present."""
    heading_index = _build_heading_index(doc)
    insert_after = None
    for idx, text, lvl in heading_index:
        if "5.1 Tools and Technologies" in text and lvl == 2:
            insert_after = idx
            break
    if insert_after is None:
        return
    anchor = doc.paragraphs[insert_after]
    p1 = _insert_paragraph_after(anchor, "5.1.6 REST API Overview", "Heading 3")
    p2 = _insert_paragraph_after(
        p1,
        "All authenticated routes use JWT Bearer tokens under the /api/v1 prefix. "
        "OpenAPI documentation is available at /docs when the API is running.",
        "Normal",
    )
    prev = p2
    for row in API_ENDPOINTS_TABLE[1:]:
        prev = _insert_paragraph_after(
            prev, f"{row[0]} {row[1]} — {row[2]}", "Normal"
        )
    _insert_paragraph_after(
        prev,
        "[INSERT SCREENSHOT: OpenAPI /docs page]",
        "Normal",
    )


def _replace_title_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs[:120]:
        t = p.text.strip()
        if not t:
            continue
        if re.search(r"project\s*404|intrusion detection system", t, re.I):
            _set_paragraph_text(p, PROJECT_NAME)
        if "Graduation" in t or "Final Year" in t or "Capstone" in t:
            if "CyberGuard" not in t:
                _set_paragraph_text(
                    p,
                    f"{PROJECT_NAME} — Technical Documentation (Graduation Project)",
                )


def update_document(input_path: Path, output_path: Path) -> dict[str, int]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)
    doc = Document(str(output_path))

    stats = {
        "heading_renames": _apply_heading_renames(doc),
        "sections_rewritten": _apply_section_bodies(doc),
        "global_replacements": _apply_global_replacements(doc),
    }
    _replace_title_paragraphs(doc)
    _update_stakeholder_table(doc)
    _update_use_case_tables(doc)
    _insert_api_section_after_tools(doc)
    stats["global_replacements"] += _apply_global_replacements(doc)

    doc.save(str(output_path))
    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description="Update CyberGuard AI DOCX documentation")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    if not args.input.is_file():
        raise SystemExit(f"Input DOCX not found: {args.input}")

    stats = update_document(args.input, args.output)
    print(f"Saved: {args.output}")
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
